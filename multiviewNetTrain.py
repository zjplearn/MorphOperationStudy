"""
author:zjp
"""


from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt  # 磅数
from docx.oxml.ns import qn  # 中文格式
from docx.shared import Inches  # 图片尺寸

from NetWithErosion import create_model
from SplitData import getTrainTest

import math
import sys
from operator import truediv

import numpy as np
from sklearn.metrics import *
import tensorflow as tf
from tensorflow.keras import optimizers
from tensorflow.keras.callbacks import ModelCheckpoint, EarlyStopping
from tensorflow.python.keras.utils.data_utils import Sequence


class DataGenerator(Sequence):

    def __init__(self, x_set, y_set, batch_size):
        """
        初始化数据生成器
        :param x_set:X shape = (tarinNum, 1, 9, 9, B//3)
        :param y_set: shape = (trainNum,1)
        :param batch_size:
        """
        self.x, self.y = x_set, y_set
        self.batch_size = batch_size

    def __len__(self):
        """
        生成器长度定义函数
        :return: 生成器的长度
        """
        return math.ceil(self.x.shape[0] / self.batch_size)

    def __getitem__(self, idx):
        """
        获取一个batch的数据
        :param idx: 当前batch的index标号
        :return: 输入数据 X, 输出数据标签
        """
        batch_x = self.x[idx * self.batch_size:(idx + 1) * self.batch_size, :, :, :, :]
        batch_y = self.y[idx * self.batch_size:(idx + 1) * self.batch_size, :]
        return batch_x, batch_y


class Trainer:

    def __init__(self, trainX, trainY, testX, testY, batchSize=16, seed=9, epoch=100, uid="", fileInfo=""):
        """
        初始化训练器
        :param batchSize:
        :param seed: 随机数种子
        :param classNum: 最终类别数量
        :param uuid: 每个用户独特的id
        """
        self.batchSize = batchSize
        self.seed = seed
        self.classNum = np.max(trainY) + 1
        self.classNum = np.max(testY) + 1
        self.trainX = trainX
        self.trainY = trainY
        self.testX = testX
        self.testY = testY
        self.epoch = epoch
        self.uid = uid
        self.fileInfo = fileInfo
        self.readData()

    def readData(self):
        """
        生成训练数据、测试数据生成器， 用于产生报告的数据
        :return:
        """
        # one_hot化标签
        TrLabel = tf.reshape(tf.one_hot(self.trainY, self.classNum), [-1, self.classNum])
        TeLabel = tf.reshape(tf.one_hot(self.testY, self.classNum), [-1, self.classNum])

        self.trainGenerator = DataGenerator(self.trainX, TrLabel, self.batchSize)

        self.testGenerator = DataGenerator(self.testX, TeLabel, self.batchSize)

        self.toReportTestX = self.testX
        print(self.toReportTestX.shape)
        self.toReportTestY = np.argmax(TeLabel, axis=1)

    def train(self):
        """
        训练以及生成报告
        :return:
        """
        net = create_model(class_num=self.classNum)
        net.compile(loss=tf.keras.losses.categorical_crossentropy, optimizer=optimizers.Adam(lr=1e-3),
                    metrics=["accuracy"])

        checkpoint = ModelCheckpoint(filepath='ModelSaved/MVN_Model/' + self.uid + self.fileInfo + '.ckpt',
                                     monitor='val_loss',
                                     mode='min',
                                     save_best_only=True,
                                     save_weights_only=True)
        early_stop = EarlyStopping(monitor='val_loss',
                                   mode='min',
                                   min_delta=0,
                                   patience=30,
                                   restore_best_weights=True)
        hist = net.fit(self.trainGenerator, epochs=self.epoch, validation_data=self.testGenerator, validation_freq=1,
                       callbacks=[early_stop, checkpoint])
        # print(net.summary())
        # 发生了错误
        # model = load_model("IPMiniGcn_Model.h5")

        # print(pred.shape)
        classification, confusion, Test_loss, Test_accuracy, oa, each_acc, aa, kappa = self.getReport(net)
        classification = str(classification)
        confusion = str(confusion)
        file = Document()

        pTitle = file.add_paragraph()
        pTitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run2 = pTitle.add_run("结果报告")
        run2.font.name = u'微软雅黑'  # 设置西文字体
        run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 设置段中文字体
        run2.font.size = Pt(21)  # 设置字体大小为21磅
        run2.font.bold = True  # 设置加粗
        pTitle.space_after = Pt(5)  # 段后距离5磅
        pTitle.space_before = Pt(5)  # 段前距离5磅

        p1 = file.add_paragraph()
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        res = ""
        res += '{} Test loss (%)'.format(Test_loss)
        res += '\n'
        res += '{} Test accuracy (%)'.format(Test_accuracy)
        res += '\n'
        res += '\n'
        res += '{} Kappa accuracy (%)'.format(kappa)
        res += '\n'
        res += '{} Overall accuracy (%)'.format(oa)
        res += '\n'
        res += '{} Average accuracy (%)'.format(aa)
        res += '\n'
        res += '\n'
        res += '{}'.format(classification)
        res += '\n'
        res += '{}'.format(confusion)
        print(res)
        run1 = p1.add_run(res)
        run1.font.name = u'仿宋_GB2312'  # 设置西文字体
        run1.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')  # 设置段中文字体
        run1.font.size = Pt(10)  # 设置字体大小为16磅
        run1.font.bold = True  # 设置加粗
        file.save( "/classification_report" + self.uid + self.fileInfo + ".docx")
        # 保存以及加载模型参数的方法
        # net.save_weights("IPMiniGcn_Model.ckpt")
        # print("save_ok")
        # net = MiniGCN(activation=self.activation, classNum=self.classNum)
        # net.compile(loss=tf.keras.losses.categorical_crossentropy, optimizer=optimizers.Adam(lr=1e-3),
        #             metrics=["accuracy"])
        # net.load_weights("IPMiniGcn_Model.ckpt")
        # print("load_ok")
        # net.evaluate(self.testGenerator)


    def getReport(self, net):
        """
        产生分类报告
        :param net: 网络模型
        :return: sklearn分类报告，混淆矩阵，测试loss, 测试精度， OA, 每类精度，AA，kappa
        """
        predLabel = np.argmax(net(self.toReportTestX), axis=1)
        classification = classification_report(self.toReportTestY, predLabel)
        oa = accuracy_score(self.toReportTestY, predLabel)
        confusion = confusion_matrix(self.toReportTestY, predLabel)
        each_acc, aa = self.aa_and_each_acc(confusion)
        kappa = cohen_kappa_score(self.toReportTestY, predLabel)
        score = net.evaluate(self.testGenerator)
        Test_Loss = score[0] * 100
        Test_accuracy = score[1] * 100

        return classification, confusion, Test_Loss, Test_accuracy, oa * 100, each_acc * 100, aa * 100, kappa * 100


    def aa_and_each_acc(self, confusion_matrix):
        """
        生成AA以及每类精度
        :param confusion_matrix:混淆矩阵
        :return: 每类精度， AA
        """
        list_diag = np.diag(confusion_matrix)
        list_raw_sum = np.sum(confusion_matrix, axis=1)
        each_acc = np.nan_to_num(truediv(list_diag, list_raw_sum))
        average_acc = np.mean(each_acc)
        return each_acc, average_acc


if __name__ == '__main__':
    trainX, trainY, testX, testY = getTrainTest(0.1)
    trainer = Trainer(trainX, trainY, testX, testY)
    trainer.train()
